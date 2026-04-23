import os
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def add_bullet(doc, bold_text, normal_text=""):
    p = doc.add_paragraph(style='List Bullet')
    r = p.add_run(bold_text)
    r.bold = True
    if normal_text:
        p.add_run(normal_text)

def create_report():
    doc = Document()

    # ── Title ──
    t = doc.add_heading('MedAssist.ai (Smart-CDSS)\nDetailed Project Report', 0)
    t.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # ── 1. Overview ──
    doc.add_heading('1. Project Overview', 1)
    doc.add_paragraph(
        "MedAssist.ai (Smart-CDSS) is an AI-powered Clinical Decision Support System that bridges "
        "the gap between patients and specialized medical care. It offers ML-driven disease prediction, "
        "LLM-powered health education, geolocation-based hospital discovery, appointment booking, "
        "and automated email reminders — all within a single, unified platform."
    )

    # ── 2. Problem Statement ──
    doc.add_heading('2. Problem Statement', 1)
    doc.add_paragraph(
        "When people feel sick, they often search online and encounter unreliable information, "
        "leading to anxiety (cyberchondria). Finding the right specialist is difficult, and patients "
        "frequently miss appointments. MedAssist.ai solves this by providing:"
    )
    add_bullet(doc, "Instant AI symptom analysis", " instead of unreliable web searches.")
    add_bullet(doc, "Curated LLM-generated health education", " to reduce misinformation.")
    add_bullet(doc, "Specialty-aware hospital routing", " connecting patients to the right doctor.")
    add_bullet(doc, "Automated appointment reminders", " (12h and 1h before) to reduce no-shows.")

    # ── 3. Market Impact ──
    doc.add_heading('3. Market Impact', 1)
    doc.add_paragraph(
        "MedAssist.ai reduces unnecessary ER visits for minor ailments while promptly routing severe "
        "conditions to appropriate specialists. For providers, it brings pre-assessed patients and "
        "cuts no-show rates via automated reminders. The global digital health market is projected "
        "to reach $550B+ by 2028, and MedAssist.ai positions itself at the intersection of AI "
        "diagnostics, telemedicine, and patient engagement."
    )

    # ── 4. How We Differ ──
    doc.add_heading('4. How Our Project Differs From Others', 1)
    doc.add_paragraph(
        "While platforms like WebMD or Practo address fragments of the healthcare journey, "
        "MedAssist.ai unifies them into one seamless workflow:"
    )
    diffs = [
        ("End-to-End Workflow: ", "Symptom input → Prediction → Education → Hospital discovery → Booking."),
        ("Calibrated ML: ", "Random Forest with CalibratedClassifierCV for realistic probability scores."),
        ("Specialty-Aware Routing: ", "Maps predicted disease to a medical specialty and prioritizes matching hospitals."),
        ("Generative AI: ", "Groq Llama-3.1 powers dynamic disease insights and an empathetic chatbot."),
        ("Voice & Accessibility: ", "Built-in Speech-to-Text for chat and Text-to-Speech for disease info."),
    ]
    for b, n in diffs:
        add_bullet(doc, b, n)

    # ── 5. Tech Stack ──
    doc.add_heading('5. Technology Stack', 1)
    stack = [
        ("Backend: ", "Python 3.x, Flask 3.0"),
        ("Database: ", "PostgreSQL on Neon DB, Flask-SQLAlchemy ORM"),
        ("Machine Learning: ", "Scikit-Learn (RandomForest, GridSearchCV, CalibratedClassifierCV), Pandas"),
        ("Frontend: ", "HTML5, CSS3, Vanilla JavaScript, Marked.js (Markdown rendering)"),
        ("Task Scheduling: ", "Flask-APScheduler (background reminder jobs)"),
        ("Deployment: ", "Render (Gunicorn WSGI), GitHub CI/CD"),
        ("Security: ", "Werkzeug password hashing, Flask session management, OTP email verification"),
    ]
    for b, n in stack:
        add_bullet(doc, b, n)

    # ── 6. APIs ──
    doc.add_heading('6. External APIs & Integrations', 1)
    apis = [
        ("Groq API (Llama-3.1-8b-instant): ", "Powers the health chatbot and generates structured JSON disease insights (severity, description, precautions, diet)."),
        ("Geoapify Places API: ", "Geolocation-based hospital search within 5 km, sorted by specialty match and proximity."),
        ("Brevo REST API: ", "Transactional emails — OTP verification, appointment confirmation, 12h/1h reminders."),
        ("Web Speech API: ", "Browser-native Speech-to-Text for voice chat input and SpeechSynthesis for reading disease info aloud."),
    ]
    for b, n in apis:
        add_bullet(doc, b, n)

    # ── 7. Workflow ──
    doc.add_heading('7. System Workflow', 1)
    steps = [
        "Registration & Login — Secure auth with email OTP verification via Brevo.",
        "Symptom Selection — User picks symptoms from a searchable grid; optionally sets age & gender.",
        "ML Prediction — Flask backend runs the calibrated Random Forest model; returns top-3 diseases with confidence %.",
        "Disease Insights — Groq LLM generates severity, description, precautions, and dietary advice as JSON.",
        "Hospital Discovery — Geoapify finds nearby facilities; system maps disease → specialty and prioritizes matches.",
        "Appointment Booking — User books a slot; record saved to PostgreSQL; confirmation email sent via Brevo.",
        "Automated Reminders — APScheduler checks every 10 min and sends 12h/1h reminder emails before appointments.",
        "AI Chatbot — Context-aware Llama-3.1 assistant with voice input support for follow-up health questions.",
    ]
    for i, s in enumerate(steps, 1):
        doc.add_paragraph(f"Step {i}. {s}", style='List Number')

    # ── 8. Inputs & Outputs ──
    doc.add_heading('8. Inputs and Outputs', 1)

    doc.add_heading('Inputs', 2)
    inputs = [
        "Patient symptoms (boolean feature vector from 130+ symptom options)",
        "Demographics — Age group (Child/Teen/Adult/Elderly), Gender (Male/Female)",
        "Geolocation — Browser GPS coordinates (latitude, longitude)",
        "Appointment data — Hospital, Doctor, Date, Time, Patient Name, Phone",
        "Chat messages — Free-text or voice-transcribed health queries",
    ]
    for i in inputs:
        doc.add_paragraph(i, style='List Bullet')

    doc.add_heading('Outputs', 2)
    outputs = [
        "Top-3 disease predictions with calibrated probability percentages",
        "Dynamic disease info — Severity, Description, Precautions, Diet (LLM-generated)",
        "Hospital list — Sorted by specialty relevance and distance, with Google Maps links",
        "Transactional emails — OTP codes, booking confirmations, timed reminders",
        "Chatbot responses — Context-aware, empathetic health guidance (Markdown-rendered)",
    ]
    for o in outputs:
        doc.add_paragraph(o, style='List Bullet')

    # ── 9. ML Methodology ──
    doc.add_heading('9. Machine Learning Methodology', 1)
    doc.add_paragraph(
        "The prediction engine uses a Random Forest Classifier trained on a medical symptoms dataset "
        "containing 130+ symptom features mapped to 40+ diseases."
    )
    ml_steps = [
        ("Data Preprocessing: ", "Column normalization, duplicate removal, handling of age/gender demographic features."),
        ("Model Training: ", "GridSearchCV over n_estimators (100-300), max_depth (20/50/None), min_samples_split (2/5), class_weight (balanced/balanced_subsample)."),
        ("Probability Calibration: ", "CalibratedClassifierCV with sigmoid method ensures prediction probabilities are realistic."),
        ("Inference: ", "Fuzzy symptom matching (difflib.get_close_matches, cutoff=0.7) handles user input variations."),
    ]
    for b, n in ml_steps:
        add_bullet(doc, b, n)

    # ── 10. Database Schema ──
    doc.add_heading('10. Database Schema', 1)
    doc.add_paragraph("PostgreSQL (Neon DB) with three core tables:")
    tables = [
        ("User: ", "id, email, password_hash, is_verified, verification_code, created_at"),
        ("PredictionRecord: ", "id, user_id (FK), symptoms (JSON), predicted_disease, top_predictions (JSON), created_at"),
        ("Appointment: ", "id, user_id (FK), hospital_name, doctor_name, appointment_date, appointment_time, patient_name, patient_phone, reminder_12h_sent, reminder_1h_sent, created_at"),
    ]
    for b, n in tables:
        add_bullet(doc, b, n)

    # ── 11. API Endpoints ──
    doc.add_heading('11. API Endpoints', 1)
    endpoints = [
        ("POST /api/register", "Register new user, send OTP email"),
        ("POST /api/verify", "Verify email with OTP code"),
        ("POST /api/login", "Authenticate user, create session"),
        ("POST /api/logout", "Destroy user session"),
        ("GET  /api/symptoms", "Return list of all available symptoms"),
        ("POST /api/predict", "Accept symptoms + demographics, return disease predictions"),
        ("GET  /api/disease-info", "Generate LLM disease details (severity, precautions, diet)"),
        ("GET  /api/hospitals", "Find nearby hospitals via Geoapify, sorted by specialty"),
        ("POST /api/book_appointment", "Save appointment to DB, send confirmation email"),
        ("GET  /api/user/data", "Retrieve user prediction history and appointments"),
        ("POST /api/chat", "Send messages to Groq LLM chatbot"),
    ]
    # Create a table
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Light Shading Accent 1'
    hdr = table.rows[0].cells
    hdr[0].text = 'Endpoint'
    hdr[1].text = 'Description'
    for ep, desc in endpoints:
        row = table.add_row().cells
        row[0].text = ep
        row[1].text = desc

    # ── 12. Conclusion ──
    doc.add_heading('12. Conclusion', 1)
    doc.add_paragraph(
        "MedAssist.ai demonstrates how modern AI technologies — from classical ML to Generative AI — "
        "can be combined into a practical, end-to-end healthcare platform. By unifying symptom analysis, "
        "disease education, hospital discovery, appointment management, and an empathetic AI chatbot, "
        "it provides a comprehensive solution that is greater than the sum of its parts. The system is "
        "production-deployed on Render with a cloud PostgreSQL backend, making it accessible to users worldwide."
    )

    # Save
    out_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'MedAssist_Project_Report.docx')
    doc.save(out_path)
    print(f"Report saved to: {out_path}")

if __name__ == "__main__":
    create_report()
